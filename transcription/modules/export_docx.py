"""Export transcript to Word DOCX format."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from config import DEFAULT_PARAGRAPH_GAP_SEC
from modules.diarize import format_segment_line
from modules.postprocess import segments_to_full_text
from utils.time_utils import seconds_to_hms


def export_docx(
    segments: list[dict[str, Any]],
    output_path: Path,
    include_timestamps: bool = True,
    *,
    group_paragraphs: bool = False,
    paragraph_gap_sec: float = DEFAULT_PARAGRAPH_GAP_SEC,
) -> Path:
    """
    Export segments to a DOCX file.

    Args:
        segments: Transcription segments.
        output_path: Destination file path.
        include_timestamps: Whether to include timestamp prefixes.
        group_paragraphs: When timestamps are off, insert blank lines on silence gaps.
        paragraph_gap_sec: Silence threshold for paragraph breaks.

    Returns:
        Path to the written file.
    """
    doc = Document()
    doc.add_heading("Meeting Transcript", level=1)

    if include_timestamps:
        for seg in segments:
            text = format_segment_line(seg).strip()
            if not text:
                continue
            start = seconds_to_hms(float(seg["start"]))
            end = seconds_to_hms(float(seg["end"]))
            doc.add_paragraph(f"[{start} -> {end}] {text}")
    else:
        body = segments_to_full_text(
            segments,
            group_paragraphs=group_paragraphs,
            paragraph_gap_sec=paragraph_gap_sec,
        )
        for paragraph in body.split("\n\n"):
            doc.add_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
